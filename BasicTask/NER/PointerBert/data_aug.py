#!/usr/bin/env python
# -*- coding: utf-8 -*-
# @Time    : 2022/10/11 11:06
# @Author  : Czq
# @File    : data_aug.py
# @Software: PyCharm

# 读取docx 文件
import json
import os
import random
import re
import pandas as pd
from docx import Document


def read_docx_file(docx_path):
    document = Document(docx_path)
    # tables = document.tables
    all_paragraphs = document.paragraphs
    return_text_list = []
    for index, paragraph in enumerate(all_paragraphs):
        one_text = paragraph.text.replace(" ", "").replace("\u3000", "")
        if one_text != "":
            return_text_list.append(one_text)
    # print(return_text_list)
    data = '\n'.join(return_text_list)
    data = data.replace('⾄', '至').replace('中华⼈民', '中华人民') \
        .replace(' ', ' ').replace(u'\xa0', ' ').replace('\r\n', '\n')
    data = re.sub("[＿_]+", "", data)
    return data


def gen_fake_by_template():
    with open('data/data_src/data_aug/template2.json', 'r', encoding='utf-8') as f:
        line = f.readlines()[0]
        template_text = json.loads(line)['text']


    W = open('data/data_src/data_aug/data_auged2.json', 'w', encoding='utf-8')
    fake_data = pd.read_csv('data/data_src/data_aug/faker_data.csv').values
    for i in range(0, len(fake_data)-200,2):
        name = str(fake_data[i][1])
        idcard = str(fake_data[i][2])
        address = str(fake_data[i][3])
        phone = str(fake_data[i][4])
        name2 = str(fake_data[i+1][1])
        idcard2 = str(fake_data[i+1][2])
        address2 = str(fake_data[i+1][3])
        phone2 = str(fake_data[i+1][4])

        # 需要注意顺序
        entities = []
        text = template_text.replace('[NAME2]', name2)
        index = text.index(name2)
        entities.append({'label': '乙方', 'start_offset': index, 'end_offset': index + len(name2)})

        text = text.replace('[ID2]', idcard2)
        index = text.index(idcard2)
        entities.append({'label': '乙方身份证号/统一社会信用代码', 'start_offset': index, 'end_offset': index + len(idcard2)})

        r = random.random()
        if r > 0.7:
            text = text.replace('（供货商）', '（供方）')
        elif 0.5 < r < 0.6:
            text = text.replace('（供货商）', '（卖方）')

        text = text.replace('[NAME1]',name)
        index = text.index(name)
        entities.append({'label':'甲方','start_offset':index,'end_offset':index+len(name)})

        text = text.replace('[ID1]', idcard)
        index = text.index(idcard)
        entities.append({'label': '甲方身份证号/统一社会信用代码', 'start_offset': index, 'end_offset': index + len(idcard)})

        text = text.replace('[NAME3]', fake_data[i + 3][1])

        text = text.replace('[CALL2]', phone2)
        index = text.index(phone2)
        entities.append({'label': '乙方联系方式', 'start_offset': index, 'end_offset': index + len(phone2)})

        text = text.replace('[ADDR2]', address2)
        index = text.index(address2)
        entities.append({'label': '乙方地址', 'start_offset': index, 'end_offset': index + len(address2)})

        text = text.replace('[NAME4]', fake_data[i + 5][1])

        text = text.replace('[CALL1]', phone)
        index = text.index(phone)
        entities.append({'label': '甲方联系方式', 'start_offset': index, 'end_offset': index + len(phone)})

        text = text.replace('[ADDR1]', address)
        index = text.index(address)
        entities.append({'label': '甲方地址', 'start_offset': index, 'end_offset': index + len(address)})

        W.write(json.dumps({'id':i, 'text':text, 'entities':entities}, ensure_ascii=False)+'\n')

    W.close()


def divided():
    import numpy as np
    file = 'data/data_src/data_aug/data_auged2.json'
    to_path = 'data/data_src/common_aug'

    with open(file, 'r', encoding='utf-8') as f:
        data = f.readlines()
    l = len(data)
    print("samples number", l)
    arr_index = list(range(l))
    random.shuffle(arr_index)

    p = int(l * 0.81)
    train_index = arr_index[:p]
    dev_index = arr_index[p:]
    print("train dataset number", len(train_index))
    print("dev dataset number", len(dev_index))
    data = np.array(data)
    train_data = data[train_index]
    dev_data = data[dev_index]

    with open(os.path.join(to_path, 'train.json'), 'w', encoding='utf-8') as f:
        for line in train_data:
            f.write(line)
    with open(os.path.join(to_path, 'dev.json'), 'w', encoding='utf-8') as f:
        for line in dev_data:
            f.write(line)


if __name__ == "__main__":
    # text = read_docx_file('data/data_src/data_aug/template2.docx')
    # print(text)
    # print(len(text))
    # with open('data/data_src/data_aug/template2.json', 'w', encoding='utf-8') as f:
    #     text = json.dumps({'text':text}, ensure_ascii=False)
    #     f.write(text)
    gen_fake_by_template()
    divided()
    pass