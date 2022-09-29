#!/usr/bin/env python
# -*- coding: utf-8 -*-
# @Time    : 2022/09/20 10:48
# @Author  : Czq
# @File    : inference.py
# @Software: PyCharm
import argparse
import random
import re
import time

import pandas as pd
from pprint import pprint

# 读取docx 文件
import torch
from docx import Document
from torch.utils.data import Dataset, DataLoader
from transformers import BertTokenizer
from DocumentReview.PointerBert.utils import set_seed, read_config_to_label
from DocumentReview.PointerBert.model_NER import PointerNERBERT


def read_docx_file(docx_path):
    document = Document(docx_path)
    # tables = document.tables
    all_paragraphs = document.paragraphs
    return_text_list = []
    for index, paragraph in enumerate(all_paragraphs):
        one_text = paragraph.text.replace(" ", "").replace("\u3000", "")
        if one_text != "":
            return_text_list.append(one_text)
    # print(return_text_list)
    data = '\n'.join(return_text_list)
    data = data.replace('⾄', '至').replace('中华⼈民', '中华人民') \
        .replace(' ', ' ').replace(u'\xa0', ' ').replace('\r\n', '\n')
    data = re.sub("[＿_]+", "", data)
    return data


def split_text(args, text):
    text_list = []
    wind = args.window_length
    step = args.window_step
    index_bias = 0
    for i in range(0,len(text),step):
        text_list.append({'index_bias':i,'text':text[i:i+wind]})
    return text_list


class TestDataset(Dataset):
    def __init__(self, data):
        self.test_data = data
    def __getitem__(self, item):
        return self.test_data[item]
    def __len__(self):
        return len(self.test_data)


def batchify_test(batch):
    sentences = []
    index_biass = []
    input_ids = []
    attention_mask = []
    token_type_ids = []
    for b in batch:
        text = b['text']
        index_bias = b['index_bias']
        sentences.append(text)
        index_biass.append(index_bias)

        input_i = [101] + tokenizer.convert_tokens_to_ids(list(text)) + [102]
        input_id = input_i.copy() + [0] * (512 - len(input_i))
        atten_mask = [1] * len(input_id) + [0] * (512 - len(input_id))
        token_type_id = [0] * 512

        assert len(input_id) == 512, len(input_id)
        input_ids.append(input_id)
        attention_mask.append(atten_mask)
        token_type_ids.append(token_type_id)

    encoded_dict = {
        'input_ids': torch.LongTensor(input_ids).to('cpu'),
        'attention_mask': torch.LongTensor(attention_mask).to('cpu'),
        'token_type_ids': torch.LongTensor(token_type_ids).to('cpu')
    }
    return encoded_dict, index_biass, sentences


def infer(args):
    set_seed(args.seed)
    # 测试输入是一个document
    # 先解析doc 或是 对text处理
    if args.input_test_file != '':
        text = read_docx_file(args.input_test_file)
    elif args.input_test_text != '':
        text = args.input_test_text
    else:
        assert False, "no input"
    print("length of text:",len(text))
    labels2id, _ = read_config_to_label(None)
    args.labels = labels2id

    # 以滑动窗口的形式切分text
    # {'index_bias':i,'text':text[i:i+wind]}
    text_list = split_text(args, text)
    dataset = TestDataset(text_list)
    dataloader = DataLoader(dataset, batch_size=args.batch_size, shuffle=False, collate_fn=batchify_test)

    model = PointerNERBERT(args).to('cpu')
    state = torch.load(args.model_load_path, map_location="cpu")
    model.load_state_dict(state['model_state'])
    # optimizer.load_state_dict(state['optimizer'])
    model.eval()
    entities = []
    for inputs, index_biass, sentences in dataloader:
        # batch_size, sentence_length, number_of_label
        start_prob, end_prob = model(inputs)
        print(start_prob.shape)
        thred = torch.FloatTensor([0.5]).to('cpu')
        start_pred = start_prob>thred
        end_pred = end_prob>thred
        # batch_size, number_of_label, sentence_length
        start_pred = start_pred.transpose(2,1)
        end_pred = end_pred.transpose(2,1)
        # if True in start_pred:
        #     print("true in start_pred")
        # if True in end_pred:
        #     print("true in end_pred")
        # 0-1 seq to entity
        for bi in range(len(start_pred)):
            index_bias = index_biass[bi]
            sentence = sentences[bi]
            for li in range(len(start_pred[bi])):
                start_seq = start_pred[bi][li]
                end_seq = end_pred[bi][li]
                start_index = []
                end_index = []
                # if True in start_seq:
                for start_ind in range(len(start_seq)):
                    if start_seq[start_ind]:
                        start_index.append(start_ind)
                        print("label:",labels2id[li], "start:", start_ind+index_bias)
                # if True in end_seq:
                for end_ind in range(len(end_seq)):
                    if end_seq[end_ind]:
                        end_index.append(end_ind)
                        print("label:", labels2id[li], "end:", end_ind+index_bias)
                # if len(start_index) == len(end_index):
                #     for start_ind, end_ind in zip(start_index, end_index):
                #         entities.append({'start':start_ind+index_bias,'end':end_ind+index_bias,'entity':sentence[start_ind:end_ind]})
                # else:
                min_len = min(len(start_index),len(end_index))
                for mi in range(min_len):
                    entities.append({'label':args.labels[li],'entity': sentence[start_index[mi]:end_index[mi]],
                        'start': start_index[mi] + index_bias, 'end': end_index[mi] + index_bias})

            # for i in range()
    pprint(entities)


if __name__ == "__main__":
    parser = argparse.ArgumentParser()

    parser.add_argument("--model_load_path", default='DocumentReview/PointerBert/model_src/PBert0926_common_all_20sche.pt')
    parser.add_argument("--input_test_file", default='', type=str, help="input file path in inference")
    parser.add_argument("--input_test_text", default='', type=str, help="input text in inference")
    # parser.add_argument("--save_path", default='ContractNER/model_src/PointerBert/pBert0920.pt')
    parser.add_argument("--window_length", default=510, type=int)
    parser.add_argument("--window_step", default=400, type=int)
    parser.add_argument("--batch_size", default=1, type=int, help="Batch size per GPU/CPU for training.")
    # parser.add_argument("--learning_rate", default=1e-5, type=float, help="The initial learning rate for Adam.")
    # parser.add_argument("--train_path", default=None, type=str, help="The path of train set.")
    # parser.add_argument("--dev_path", default=None, type=str, help="The path of dev set.")
    parser.add_argument("--max_seq_len", default=512, type=int, help="The maximum input sequence length. "
                                                                     "Sequences longer than this will be split automatically.")
    parser.add_argument("--bert_emb_size", default=768, type=int, help="The embedding size of pretrained model")
    parser.add_argument("--hidden_size", default=200, type=int, help="The hidden size of model")
    parser.add_argument("--num_epochs", default=100, type=int, help="Total number of training epochs to perform.")
    parser.add_argument("--seed", default=1000, type=int, help="Random seed for initialization")
    parser.add_argument("--logging_steps", default=200, type=int, help="The interval steps to logging.")
    parser.add_argument("--valid_steps", default=100, type=int,
                        help="The interval steps to evaluate model performance.")
    parser.add_argument('--device', choices=['cpu', 'cuda'], default="cuda",
                        help="Select which device to train model, defaults to gpu.")
    # parser.add_argument("--model", choices=["uie-base", "uie-tiny", "uie-medium", "uie-mini", "uie-micro", "uie-nano"],
    #                     default="uie-base", type=str, help="Select the pretrained model for few-shot learning.")
    parser.add_argument("--model", default="model/language_model/chinese-roberta-wwm-ext", type=str,
                        help="Select the pretrained model for few-shot learning.")
    parser.add_argument("--init_from_ckpt", default=None, type=str,
                        help="The path of model parameters for initialization.")
    args = parser.parse_args()

    pprint(args)
    tokenizer = BertTokenizer.from_pretrained('model/language_model/chinese-roberta-wwm-ext')
    # args.input_test_text = """个人借条今胡兵和王英由于夫妻日常生活开销向甘红霞借款贰万元整（人民币20000），借款日期为实际收到借款的当天，出借人甘红霞直接微信向胡兵转账20000元，借款期限为一年，年利率为13%，如不能按时归还，和愿承担产生的一切法律责任。特以此为据借款人1签字：身份证号：542123197402087786借款人2签字：身份证号：220781193602287047借条日期：2022.7.20"""
    # print(len(args.input_test_text))
    # args.input_test_text = """借条借款人姓名：孙林身份证号：孙林因生活所需，本人向叶芳（出借人）借到现金人民币（大写）叁仟伍佰圆整（￥3500元），且已经收到出借人支付的上述款项。签署时间：2022年2月1日借款人签字：借款人联系方式：15125822509保证人签名：保证人身份证号：511503197708235948保证人联系方式：13820935531"""
    # print(len(args.input_test_text))
    # args.input_test_text = "\n设备租赁合同5出租⽅：珑珑租赁公司(简称甲⽅)承租⽅：花鸟有限公司(简称⼄⽅)甲、⼄双⽅根据《中国⼯商银⾏X市信托部设备租赁业务试⾏办法》的规定，签订设备租赁合同，并商定如下条款，共同遵守执⾏。⼀、甲⽅根据⼄⽅上级批准的项⽬和⼄⽅⾃⾏选定的设备和技术质量标准，购进以下设备租给⼄⽅使⽤。⼆、甲⽅根据与⽣产⼚(商)签订的设备订货合同规定，于2022年春季交货，由供货单位直接发运给⼄⽅。⼄⽅直接到供货单位⾃提⾃运。⼄⽅收货后应⽴即向甲⽅开回设备收据。三、设备的验收、安装、调试、使⽤、保养、维修管理等，均由⼄⽅⾃⾏负责。设备的质量问题由⽣产⼚负责，并在订货合同中予以说明。四、设备在租赁期间的所有权属于甲⽅。⼄⽅收货后，应以甲⽅名义向当地保险公司投保综合险，保险费由⼄⽅负责。⼄⽅应将投保合同交甲⽅作为本合同附件。五、在租赁期间，⼄⽅享有设备的使⽤权，但不得转让或作为财产抵押，未经甲⽅同意亦不得在设备上增加或拆除任何部件和迁移安装地点。甲⽅有权检查设备的使⽤和完好情况，⼄⽅应提供⼀切⽅便。六、设备租赁期限为2年，租期从供货⼚向甲⽅托收货款时算起，租⾦总额为⼈民币24240元(包括⼿续费1%)，分6期交付，每期租金4040元，由甲⽅在每期期末按期向⼄⽅托收。如⼄⽅不能按期承付租⾦，甲⽅则按逾期租⾦总额每天加收万分之三的罚⾦。七、本合同⼀经签订不能撤销。如⼄⽅提前交清租⾦，结束合同，甲⽅给予退还⼀部分利息的优惠。⼋、本合同期满，甲⽅同意按⼈民币20000元的优惠价格将设备所有权转给⼄⽅。九、⼄⽅上级单位同意作为⼄⽅的经济担保⼈，负责⼄⽅切实履⾏本合同各条款规定，如⼄⽅在合同期内不能承担合同中规定的经济责任时，担保⼈应向甲⽅⽀付⼄⽅余下的各期租⾦和其他损失。⼗、本合同经双⽅和⼄⽅担保⼈盖章后⽣效。本合同正本两份，甲、⼄⽅各执⼀份;副本两份份，⼄⽅担保⼈和⼄⽅开户银⾏各⼀份。甲⽅：珑珑租赁公司(签章)⼄⽅：花鸟有限公司(公章)2022年1⽉1⽇\n"
    # print(len(args.input_test_text))
    start_time = time.time()
    args.input_test_file = 'data/DocData/laodong/ld1.docx'
    infer(args)
    print('used time:', time.time()-start_time)
