#!/usr/bin/env python
# -*- coding: utf-8 -*-
# @Time    : 2022/09/15 16:26
# @Author  : Czq
# @File    : read_docx_server.py
# @Software: PyCharm
import json
import os
import re
import uuid
import requests
from docx import Document
from flask import Flask, flash, request, redirect, url_for
from werkzeug.utils import secure_filename

UPLOAD_FOLDER = 'data/uploads/'
ALLOWED_EXTENSIONS = {'txt', 'docx'}

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
# 16M
app.config['MAX_CONTENT_LENGTH'] = 16 * 1000 * 1000
app.secret_key = 'qweasdqwe'

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/file_link_path_to_text', methods=['GET', 'POST'])
def file_link_path_to_text():
    if request.method == 'POST':
        in_json = request.get_data()
        if in_json is not None:
            # json_data = json.loads(in_json.decode("utf-8"))
            json_data = request.get_json()
            print(json_data)
            file_path = json_data.get('file_path','')
            if file_path == '':
                return json.dumps({'text': '', 'status': 0, 'message': 'no file'}, ensure_ascii=False)
            if not ('.docx' in file_path or '.txt' in file_path):
                return json.dumps({'text': '', 'status': 0, 'message': 'not allowed file, only .docx and .txt'}, ensure_ascii=False)
            os.system('cd data/uploads && wget '+file_path)
            filename = file_path.split('/')[-1]
            if not os.path.exists(os.path.join('data/uploads',filename)):
                return json.dumps({'text': '', 'status': 0, 'message': 'not get file'}, ensure_ascii=False)
            if '.docx' in file_path:
                data = read_docx_file(os.path.join('data/uploads',filename))
            elif '.txt' in file_path:
                data = read_txt_file(os.path.join('data/uploads', filename))
            else:
                 return json.dumps({'text': '', 'status': 0, 'message': 'not allowed file, only .docx and .txt'},
                                  ensure_ascii=False)
            os.remove(os.path.join('data/uploads',filename))
            return json.dumps({'text': data, 'status':1, 'message':'ok'},ensure_ascii=False)
        else:
            return json.dumps({'text': '', 'status': 0, 'message': 'no input'}, ensure_ascii=False)
    else:
        return json.dumps({'text': '', 'status': 0, 'message': 'please use post method'}, ensure_ascii=False)




@app.route('/upload_docx_to_get_text', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'POST':
        # check if the post request has the file part
        if 'file' not in request.files:
            flash('No file part')
            return json.dumps({'text':'', 'status': 0,'message':'error, no file part'},ensure_ascii=False)
        file = request.files['file']
        # If the user does not select a file, the browser submits an
        # empty file without a filename.
        if file.filename == '':
            flash('No selected file')
            return json.dumps({'text': '', 'status': 0, 'message': 'No selected file'},ensure_ascii=False)
        if file and allowed_file(file.filename):
            # filename = secure_filename(file.filename)
            filename = file.filename
            if '.' in filename:
                if filename.rsplit('.', 1)[1].lower() == 'txt':
                    t = uuid.uuid3(uuid.NAMESPACE_DNS, filename)
                    filename = str(t) + '.txt'
                    file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                    data = read_txt_file(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                    os.remove(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                elif filename.rsplit('.', 1)[1].lower() == 'docx':
                    t = uuid.uuid3(uuid.NAMESPACE_DNS, filename)
                    filename = str(t) + '.docx'
                    file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                    data = read_docx_file(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                    os.remove(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                else:
                    return json.dumps({'text':'','status':0,'message':'error, not allowed file, must use .docx or .txt'},ensure_ascii=False)
                return json.dumps({'text': data, 'status':1, 'message':'ok'},ensure_ascii=False)
            else:
                return json.dumps({'text': '', 'status':0,'message': 'error, not allowed file or no select file'},ensure_ascii=False)

    return json.dumps({'text': '', 'status': 0, 'message': 'not post'},ensure_ascii=False)


# 读取docx 文件
def read_docx_file(docx_path):
    document = Document(docx_path)
    # tables = document.tables
    all_paragraphs = document.paragraphs
    return_text_list = []
    for index, paragraph in enumerate(all_paragraphs):
        one_text = paragraph.text.replace(" ", "").replace("\u3000", "")
        if one_text != "":
            return_text_list.append(one_text)
    # print(return_text_list)
    data = '\n'.join(return_text_list)
    data = data.replace('⾄', '至').replace('中华⼈民', '中华人民') \
        .replace(' ', '').replace(u'\xa0', '').replace('\r\n', '\n')
    data = re.sub("[＿_]+", "", data)
    return data

def read_txt_file(txt_path):
    return_text_list = []
    with open(txt_path,'r',encoding='utf-8')as f:
        for line in f.readlines():
            return_text_list.append(line.strip())
    data = '\n'.join(return_text_list)
    data = data.replace('⾄', '至').replace('中华⼈民', '中华人民') \
        .replace(' ', '').replace(u'\xa0', '').replace('\r\n', '\n')
    data = re.sub("[＿_]+", "", data)
    return data


if __name__ == '__main__':
    app.run(host="0.0.0.0", port=8111, debug=False)  # , use_reloader=False)
    # app.run(host="101.69.229.138", port=8111, debug=False)  # , use_reloader=False)
