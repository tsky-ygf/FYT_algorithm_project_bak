#!/usr/bin/env python
# -*- coding: utf-8 -*-
# @Time    : 2022/3/17 10:21
# @Author  : Adolf
# @Site    : 
# @File    : parse_word.py
# @Software: PyCharm
from docx import Document


# def get_text_from_docx(path):
#     """
#     docx文件 提取 文本内容, 不做任何的加工处理
#     """
#     document = Document(path)
#     # 读取每段资料
#     texts = [paragraph.text for paragraph in document.paragraphs]
#     return texts


# 读取docx 文件
def read_docx_file(docx_path):
    document = Document(docx_path)
    # tables = document.tables
    # for table in tables:
    # print(table)
    # print(len(table.rows))
    # print(len(table.columns))
    # print(table.cell(6, 5).text)
    # for i in range(len(table.rows)):
    #     print(table.rows[i])

    all_paragraphs = document.paragraphs

    return_text_list = []

    for index, paragraph in enumerate(all_paragraphs):
        one_text = paragraph.text.replace(" ", "").replace("\u3000", "")
        if one_text != "":
            return_text_list.append(one_text)
    # print(return_text_list)
    return return_text_list


if __name__ == '__main__':
    res = read_docx_file(docx_path="data/DocData/maimai/test.docx")
    # print(res)
