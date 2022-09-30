#!/usr/bin/env python
# -*- coding: utf-8 -*-
# @Time    : 2022/7/21 13:23
# @Author  : Adolf
# @Site    : 
# @File    : contract_st_show.py
# @Software: PyCharm
import os
import re

import requests as requests
import streamlit as st
from annotated_text import annotated_text
from docx import Document

# from paddlenlp import Taskflow
# import pycorrector
from loguru import logger
from pprint import pprint
from pypinyin import pinyin, lazy_pinyin

# from annotated_text import annotated_text

# text_correction = Taskflow("text_correction")

# 读取docx 文件
def read_docx_file(docx_path):
    document = Document(docx_path)
    # tables = document.tables
    all_paragraphs = document.paragraphs
    return_text_list = []
    for index, paragraph in enumerate(all_paragraphs):
        one_text = paragraph.text.replace(" ", "").replace("\u3000", "")
        if one_text != "":
            return_text_list.append(one_text)
    # print(return_text_list)
    data = '\n'.join(return_text_list)
    data = data.replace('⾄', '至').replace('中华⼈民', '中华人民') \
        .replace(' ', ' ').replace(u'\xa0', ' ').replace('\r\n', '\n')
    data = re.sub("[＿_]+", "", data)
    return data

@st.cache
def get_data(_file):
    _text = read_docx_file(_file)
    return _text


contract_type = st.sidebar.selectbox("请选择合同类型",
                                     ["借条", "借款", "劳动", '房屋租赁', '买卖', '劳务', '保密', '采购', '一般租赁'],
                                     key="合同类型")
mode_type = st.sidebar.selectbox("请选择上传数据格式", ["docx", "文本", "txt"], key="text")
usr = st.sidebar.selectbox("请选择立场", ['甲方', '乙方'], key="中立方")
is_show = st.sidebar.selectbox("请选择是否用于对外展示", ['是', '否'], key="show")

if is_show == "是":
    is_show = True
else:
    is_show = False

contract_type = ''.join(lazy_pinyin(contract_type))
config_path = "DocumentReview/Config/{}.csv".format(contract_type)
# model_path = "model/uie_model/new/{}/model_best/".format(contract_type)
model_path = "model/uie_model/export_cpu/{}/inference".format(contract_type)
# print(contract_type)

if usr == '甲方':
    usr = 'Part A'
    usr2 = "party_a"
else:
    usr = 'Part B'
    usr2 = "party_b"

if mode_type == "docx":
    file = st.file_uploader('上传文件', type=['docx'], key=None)
    text = get_data(file)
elif mode_type == "txt":
    text = st.text_area(label='请输入文本', value='', height=600, key=None)
elif mode_type == "文本":
    text = st.text_area(label='请输入文本', value='', height=600, key=None)
else:
    raise Exception("暂时不支持该数据格式")

correct = st.button("文本纠错")
run = st.button("开始审核")

if is_show:
    from DocumentReview.ContractReview.showing_sample import BasicUIEAcknowledgementShow
    # acknowledgement = BasicUIEAcknowledgementShow(config_path=config_path,
    #                                               model_path=model_path,
    #                                               device="cpu", )
    # print(resp_json)
    pass
else:
    from DocumentReview.ContractReview.basic_contract import BasicUIEAcknowledgement

    acknowledgement = BasicUIEAcknowledgement(config_path=config_path,
                                              model_path=model_path,
                                              device="cpu", )
# result = requests.post(url,json={"user_standpoint_id":usr})


# use paddleNLP to correct the text
# if correct:
#     res_correct = text_correction(text)
#     if len(res_correct) > 0:
#         # st.write(annotated_text(res_correct))
#         # st.write("纠错后的文本")
#         # st.write(text)
#         print(res_correct)
#         res_text = []
#         last_index = 0
#         for one_error in res_correct[0]['errors']:
#             one_position = one_error['position']
#             print(one_position)
#             res_text.append(text[last_index:one_position])
#             res_text.append((text[one_position], one_error['correction'][text[one_position]], '#FF8B72'))
#             # res_text.append(text[one_position + 1:])
#             last_index = one_position + 1
#         res_text.append(text[last_index:])
#         annotated_text(*res_text)
#     else:
#         st.write("错别字审核通过")
if correct:
    import requests
    import pandas as pd

    # st.write(text)
    r = requests.post("http://172.19.82.199:6598/macbert_correct", json={"text": text})
    result = r.json()
    if result['success']:
        result = result['result']
        # if len(result) > 0:
        #     st.write(result)
        # st.write(result)
        origin_text_list = []
        correct_text_list = []
        error_list = []
        for one_error in result:
            if len(one_error[2]) == 0:
                # st.write("{} ====> 纠错通过".format(one_error[0]))
                pass
            else:
                # st.write("{} ====> {}".format(one_error[0], one_error[1]))
                # res_dict[one_error[0]] = one_error[1]
                origin_text_list.append(one_error[0])
                correct_text_list.append(one_error[1])
                error_text = ""
                # error_list.append(one_error[2][0])
                for er in one_error[2]:
                    error_text += er[0] + "===>" + er[1] + "\n"
                st.write(error_text)
                error_list.append(error_text)
                # for
        res_dict = {"原始文本": origin_text_list, "纠错后的文本": correct_text_list, "错别字": error_list}
        # print(res_dict)

        my_df = pd.DataFrame.from_dict(res_dict)
        st.table(my_df)
        # my_df.columns = ['纠错后的句子']
        # st.dataframe(my_df)

    else:
        logger.error(result['error_msg'])
        # result = []

if run:
    # url = "http://101.69.229.138:8111/upload_docx_to_get_text"
    # response_text = requests.post(url, files={'file':file.getvalue()})
    # print("response_text_________________________")
    # print(response_text.text)
    url = "http://172.19.82.199:8110/get_contract_review_result"
    req_data = {
        "contract_type_id": contract_type,
        "user_standpoint_id": usr2,
        "contract_content": text
    }
    resp_json = requests.post(url, json=req_data).json()
    # acknowledgement.review_main(content=text, mode="text", usr=usr)
    # pprint(acknowledgement.review_result, sort_dicts=False)
    # st.write(resp_json)
    index = 1

    # st.write(acknowledgement.data)
    # origin_text = acknowledgement.data
    origin_text = text
    origin_text = list(origin_text)
    length_origin_text = len(origin_text)
    value_list = resp_json['result']
    for value_en in value_list:
        key = value_en.get('review_point','')
        value = {}
        value['审核结果'] = value_en.get('review_result','')
        value['内容'] = value_en.get('review_content','')
        value['风险等级'] = value_en.get('risk_level','')
        value['风险点'] = value_en.get('risk_point','')
        value['法律建议'] = value_en.get('legal_advice','')
        value['法律依据'] = value_en.get('legal_basis','')
        value['start'] = value_en.get('review_content_start','')
        value['end'] = value_en.get('review_content_end','')

    # for key, value in acknowledgement.review_result.items():
        # st.write(key, value)
        # st.wr(value)
        st.markdown('### {}、审核点：{}'.format(index, key))
        index += 1
        try:
            if "审核结果" in value and value["审核结果"] != "":
                st.markdown("审核结果：{}".format(value['审核结果']))

            if value['审核结果'] == "通过" and value["风险等级"] == "低":
                continue

            if "内容" in value and value["内容"] != "":
                st.markdown("审核内容：{}".format(value['内容']))
            if "法律建议" in value and value["法律建议"] != "":
                st.markdown("法律建议：{}".format(value['法律建议']))
            if "风险点" in value and value["风险点"] != "":
                st.markdown("风险点：{}".format(value['风险点']))

            if "法律依据" in value and value["法律依据"] != "":
                st.markdown("法律依据：{}".format(value['法律依据']))
            dang = 0
            if "风险等级" in value and value["风险等级"] != "":
                st.markdown("风险等级：{}".format(value['风险等级']))
                if value['风险等级'] == '低':
                    dang = "#8ef"
                elif value['风险等级'] == '中':
                    dang = "#afa"
                elif value['风险等级'] == '高':
                    dang = "#faa"

            if 'start' in value and value['start'] != "" and 'end' in value and value['end'] != "":
                st.markdown("start: {}".format(value['start']))
                st.markdown("end: {}".format(value['end']))
                starts = str(value['start']).split('#')
                starts = list(filter(None, starts))
                starts = list(map(lambda x: int(x), starts))
                ends = str(value['end']).split('#')
                ends = list(filter(None, ends))
                ends = list(map(lambda x: int(x), ends))
                if dang != 0:
                    for start, end in zip(starts, ends):
                        for ind in range(start, end):
                            if isinstance(origin_text[ind], str):
                                origin_text[ind] = (str(origin_text[ind]), dang)
                            # st.write(origin_text[ind])
        except Exception as e:
            print('-'*50+'error!')
            print(e)
            st.write("这个审核点小朱正在赶制，客官请稍等。。。。可爱")
    for i in range(len(origin_text)):
        if not isinstance(origin_text[i], str):
            origin_text[i] = (str(origin_text[i][0]), '', str(origin_text[i][1]))
    annotated_text(*origin_text)
    st.write('-'*100)
    st.text(text)
