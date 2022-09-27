#!/usr/bin/env python
# -*- coding: utf-8 -*-
# @Time    : 2022/9/26 09:43
# @Author  : Adolf
# @Site    : 
# @File    : contract_for_server.py
# @Software: PyCharm
import json
import os
import re
import uuid

from docx import Document

from DocumentReview.ContractReview.showing_sample import BasicUIEAcknowledgement

CONTRACT_SERVER_DATA_PATH = "DocumentReview/Config/contract_server_data.json"


def get_support_contract_types():
    with open(CONTRACT_SERVER_DATA_PATH, "r", encoding="utf-8") as f:
        return json.load(f).get("support_contract_types")


def get_user_standpoint():
    with open(CONTRACT_SERVER_DATA_PATH, "r", encoding="utf-8") as f:
        user_standpoints = json.load(f).get("user_standpoints")
    return user_standpoints


def get_contract_type_list():
    support_contract_types = get_support_contract_types()
    return [item.get("type_id") for item in support_contract_types]


def init_model():
    contract_type_list = get_contract_type_list()
    acknowledgement_dict = {}
    for contract_type in contract_type_list:
        config_path = "DocumentReview/Config/{}.csv".format(contract_type)
        model_path = "model/uie_model/export_cpu/{}/inference".format(contract_type)
        acknowledgement_dict[contract_type] = BasicUIEAcknowledgement(config_path=config_path,
                                                                      model_path=model_path,
                                                                      device="cpu",
                                                                      logger_file='log/contract_review/model.log')
    return acknowledgement_dict


def get_text_from_file_link_path(file_link):
    os.system('cd data/uploads && wget ' + file_link)
    filename = file_link.split('/')[-1]
    if '.docx' in file_link:
        data = read_docx_file(os.path.join('data/uploads', filename))
    elif '.txt' in file_link:
        data = read_txt_file(os.path.join('data/uploads', filename))
    else:
        data =  "invalid input"
    os.remove(os.path.join('data/uploads', filename))
    return data


# 读取docx 文件
def read_docx_file(docx_path):
    document = Document(docx_path)
    # tables = document.tables
    all_paragraphs = document.paragraphs
    return_text_list = []
    for index, paragraph in enumerate(all_paragraphs):
        one_text = paragraph.text.replace(" ", "").replace("\u3000", "")
        if one_text != "":
            return_text_list.append(one_text)
    # print(return_text_list)
    data = '\n'.join(return_text_list)
    data = data.replace('⾄', '至').replace('中华⼈民', '中华人民') \
        .replace(' ', ' ').replace(u'\xa0', ' ').replace('\r\n', '\n')
    data = re.sub("[＿_]+", "", data)
    return data

def read_txt_file(txt_path):
    return_text_list = []
    with open(txt_path,'r',encoding='utf-8')as f:
        for line in f.readlines():
            return_text_list.append(line.strip())
    data = '\n'.join(return_text_list)
    data = data.replace('⾄', '至').replace('中华⼈民', '中华人民') \
        .replace(' ', ' ').replace(u'\xa0', ' ').replace('\r\n', '\n')
    data = re.sub("[＿_]+", "", data)
    return data


def get_text_from_file(file):
    filename = file.filename
    if filename.rsplit('.', 1)[1].lower() == 'txt':
        t = uuid.uuid3(uuid.NAMESPACE_DNS, filename)
        filename = str(t) + '.txt'
        file.save(os.path.join('data/uploads', filename))
        data = read_txt_file(os.path.join('data/uploads', filename))
        os.remove(os.path.join('data/uploads', filename))
    elif filename.rsplit('.', 1)[1].lower() == 'docx':
        t = uuid.uuid3(uuid.NAMESPACE_DNS, filename)
        filename = str(t) + '.docx'
        file.save(os.path.join('data/uploads', filename))
        data = read_docx_file(os.path.join('data/uploads', filename))
        os.remove(os.path.join('data/uploads', filename))
    else:
        data = "invalid input"

    return data


