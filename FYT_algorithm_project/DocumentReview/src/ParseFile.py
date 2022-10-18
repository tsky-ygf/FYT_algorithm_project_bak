#!/usr/bin/env python
# -*- coding: utf-8 -*-
# @Time    : 2022/10/09 17:31
# @Author  : Czq
# @File    : ParseFile.py
# @Software: PyCharm
import re

from docx import Document


def read_txt_file(txt_path):
    return_text_list = []
    with open(txt_path,'r',encoding='utf-8')as f:
        for line in f.readlines():
            return_text_list.append(line.strip())
    data = '\n'.join(return_text_list)
    data = data.replace('⾄', '至').replace('中华⼈民', '中华人民') \
        .replace(' ', ' ').replace(u'\xa0', ' ').replace('\r\n', '\n')
    data = re.sub("[＿_]+", "", data)
    return data




# 读取docx 文件
def read_docx_file(docx_path):
    document = Document(docx_path)
    # tables = document.tables
    all_paragraphs = document.paragraphs
    return_text_list = []
    for index, paragraph in enumerate(all_paragraphs):
        one_text = paragraph.text.replace(" ", "").replace("\u3000", "")
        if one_text != "":
            return_text_list.append(one_text)
    # print(return_text_list)
    data = '\n'.join(return_text_list)
    data = data.replace('⾄', '至').replace('中华⼈民', '中华人民') \
        .replace(' ', ' ').replace(u'\xa0', ' ').replace('\r\n', '\n')
    data = re.sub("[＿_]+", "", data)
    return data

