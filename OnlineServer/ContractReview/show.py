# !/usr/bin/env python
# -*- coding: utf-8 -*-
# @Time    : 2022/7/21 13:23
# @Author  : Adolf
# @Site    :
# @File    : contract_st_show.py
# @Software: PyCharm
import re

import streamlit as st
from annotated_text import annotated_text
from docx import Document

from loguru import logger


# 读取docx 文件
def read_docx_file(docx_path):
    document = Document(docx_path)
    # tables = document.tables
    all_paragraphs = document.paragraphs
    return_text_list = []
    for index, paragraph in enumerate(all_paragraphs):
        one_text = paragraph.text.replace(" ", "").replace("\u3000", "")
        if one_text != "":
            return_text_list.append(one_text)
    # print(return_text_list)
    data = '\n'.join(return_text_list)
    data = data.replace('⾄', '至').replace('中华⼈民', '中华人民') \
        .replace(' ', '').replace(u'\xa0', '').replace('\r\n', '\n')
    data = re.sub("[＿_]+", "", data)
    return data


@st.cache
def get_data(_file):
    _text = read_docx_file(_file)
    return _text


def contract_review_main():
    import requests
    contract_type_list = requests.get("http://127.0.0.1:8131/get_contract_type").json()["result"]

    contract_mapping = {one["contract_type"]: one["type_id"] for one in contract_type_list}

    contract_type = st.sidebar.selectbox("请选择合同类型",
                                         [con_type["contract_type"] for con_type in contract_type_list],
                                         key="合同类型")
    mode_type = st.sidebar.selectbox("请选择上传数据格式", ["docx", "文本", "txt"], key="text")

    usr_list = requests.get("http://127.0.0.1:8131/get_user_standpoint").json()["result"]
    usr_mapping = {one["standpoint"]: one["id"] for one in usr_list}

    usr = st.sidebar.selectbox("请选择立场", [usr["standpoint"] for usr in usr_list], key="中立方")
    # is_show = st.sidebar.selectbox("请选择是否用于对外展示", ['是', '否'], key="show")

    # if is_show == "是":
    #     is_show = True
    # else:
    #     is_show = False

    if mode_type == "docx":
        file = st.file_uploader('上传文件', type=['docx'], key=None)
        text = get_data(file)
    elif mode_type == "txt":
        text = st.text_area(label='请输入文本', value='', height=600, key=None)
    elif mode_type == "文本":
        text = st.text_area(label='请输入文本', value='', height=600, key=None)
    else:
        raise Exception("暂时不支持该数据格式")

    correct = st.button("文本纠错")
    run = st.button("开始审核")

    if correct:
        import requests
        from OnlineServer.ContractCorrector.show import corrector_main
        st.write("文本纠错start")
        r = requests.post("http://127.0.0.1:6598/get_corrected_contract_result", json={"text": text})
        result = r.json()
        st.write(result)
        corrector_main(text, result)

    if run:
        url = "http://127.0.0.1:8131/get_contract_review_result"
        req_data = {
            "contract_type_id": contract_mapping[contract_type],
            "user_standpoint_id": usr_mapping[usr],
            "contract_content": text
        }
        resp_json = requests.post(url, json=req_data).json()["result"]
        index = 1

        origin_text = text
        origin_text = list(origin_text)
        value_list = resp_json
        for value_en in value_list:
            key = value_en.get('review_point', '')
            value = dict()
            value['审核结果'] = value_en.get('review_result', '')
            value['内容'] = value_en.get('review_content', '')
            value['风险等级'] = value_en.get('risk_level', '')
            value['风险点'] = value_en.get('risk_point', '')
            value['法律建议'] = value_en.get('legal_advice', '')
            value['法律依据'] = value_en.get('legal_basis', '')
            value['start'] = value_en.get('review_content_start', '')
            value['end'] = value_en.get('review_content_end', '')

            st.markdown('### {}、审核点：{}'.format(index, key))
            index += 1
            try:
                if "审核结果" in value and value["审核结果"] != "":
                    st.markdown("审核结果：{}".format(value['审核结果']))

                if value['审核结果'] == "通过" and value["风险等级"] == "低":
                    continue

                if "内容" in value and value["内容"] != "":
                    st.markdown("审核内容：{}".format(value['内容']))
                if "法律建议" in value and value["法律建议"] != "":
                    st.markdown("法律建议：{}".format(value['法律建议']))
                if "风险点" in value and value["风险点"] != "":
                    st.markdown("风险点：{}".format(value['风险点']))

                if "法律依据" in value and value["法律依据"] != "":
                    st.markdown("法律依据：{}".format(value['法律依据']))
                dang = 0
                if "风险等级" in value and value["风险等级"] != "":
                    st.markdown("风险等级：{}".format(value['风险等级']))
                    if value['风险等级'] == '低':
                        dang = "#8ef"
                    elif value['风险等级'] == '中':
                        dang = "#afa"
                    elif value['风险等级'] == '高':
                        dang = "#faa"

                if 'start' in value and value['start'] != "" and 'end' in value and value['end'] != "":
                    st.markdown("start: {}".format(value['start']))
                    st.markdown("end: {}".format(value['end']))
                    starts = str(value['start']).split('#')
                    starts = list(filter(None, starts))
                    starts = list(map(lambda x: int(x), starts))
                    ends = str(value['end']).split('#')
                    ends = list(filter(None, ends))
                    ends = list(map(lambda x: int(x), ends))
                    if dang != 0:
                        for start, end in zip(starts, ends):
                            for ind in range(start, end):
                                if isinstance(origin_text[ind], str):
                                    origin_text[ind] = (str(origin_text[ind]), dang)
                                # st.write(origin_text[ind])
            except Exception as e:
                print('-' * 50 + 'error!')
                print(e)
                st.write("这个审核点小陈正在赶制，客官请稍等。。。。可爱")
        for i in range(len(origin_text)):
            if not isinstance(origin_text[i], str):
                origin_text[i] = (str(origin_text[i][0]), '', str(origin_text[i][1]))
        annotated_text(*origin_text)
        st.write('-' * 100)
        st.text(text)
